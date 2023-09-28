import os
import configparser
nameConfFile = 'PywordData.ini'
config = configparser.ConfigParser()
if nameConfFile in os.listdir():
   config.read(nameConfFile)
   initialdir = config['Main']['savepath']
else:
   open(nameConfFile, 'x')
   initialdir = os.getenv('USERPROFILE')+"\\Desktop"
   config["Main"]={}
   config["Main"]['savepath'] = initialdir

import tkinter
import tkinter.filedialog
rootWindow = tkinter.Tk()
rootWindow.wm_withdraw()#для скрытия окна
wordFileName = tkinter.filedialog.askopenfilename(initialdir = initialdir, title = "Выберите файл с РП", filetypes = [('Только word файлы', '.docx .doc')])
if wordFileName:
   config['Main']['savepath'] = "/".join(wordFileName.split("/")[0:-1])
   with open(nameConfFile, 'w') as configFile:
      config.write(configFile)
   
   import openpyxl
   from openpyxl.styles import PatternFill
   wb = openpyxl.Workbook()
   sheet = wb['Sheet']
   sheet.title = "Лист1"
   sheet['A1'] = 'Глава/Раздел'
   sheet['B1'] = 'Тема занятия'
   sheet['C1'] = 'Домашнее задание'
   sheet['D1'] = 'Тип занятия'
   sheet['E1'] = 'Количество часов'
   
   import docx
   import re
   doc = docx.Document(wordFileName)
   def newCellExcel():
      global sheet, excelI
      sheet.cell(row=excelI, column=1).value='i'
      sheet.cell(row=excelI, column=3).value='.'
      sheet.cell(row=excelI, column=4).value='обычный'
      sheet.cell(row=excelI, column=5).value=1
      excelI+=1
   def findColumn(nameOfColumn, row): 
      for i in row:
         if nameOfColumn in i.text:
            return row.index(i)
      while True:
         print("Программа не нашла колонки с названием \"" + nameOfColumn + "\". Исправьте документ и запустите программу повторно или укажите фразу (большие буквы важны!), по которой надо искать колонку, ниже:")
         nameOfColumn = input()
         for i in row:
            if nameOfColumn in i.text:
               return row.index(i)
   
   def calculateWidthOfTheme(row, columnNum):
      global widthOfTheme
      widthOfTheme=1
      while columnNum+widthOfTheme<len(row) and row[columnNum+widthOfTheme]._tc==row[columnNum]._tc:
         widthOfTheme+=1
   
   def appendToExcelWithoutRepeat(whatTextAppend):
      global sheet, excelI
      sheet.cell(row=excelI, column=2).value=str(excelI-1)+' '+' '.join(whatTextAppend.split())
      newCellExcel()
   
   def appendToExcelWithMerge(newText, oldText):
      global docTable, excelI
      if ''.join(newText.split()) != ''.join(oldText.split()):#тексты не равны
         excelCell = sheet.cell(row = excelI-1, column=2)
         if excelCell.value[-1] != '.':
            excelCell.value = excelCell.value+'.'
         excelCell.value = excelCell.value+' '+" ".join(newText.split())

   def correctThemeHorizontal(thisCell):
      global docTable, columnNum, i
      if thisCell._tc==docTable.column_cells(columnNum-1)[i]._tc:
         print("Обнаружен возможный сдвиг таблицы. Обратите внимание на текст в последней выведенной на экран ячейки. Если он является темой урока - нажмите 1. Иначе - 2")
         if (int(input()) == 2):
            while docTable.column_cells(columnNum)[i]._tc==docTable.column_cells(columnNum-1)[i]._tc:
               columnNum+=1 #корректировка тем по горизонтали (бывает уезжает)
         calculateWidthOfTheme(tempRow, columnNum)
      return docTable.column_cells(columnNum)[i]
   
   def head():
      global i, mainI, isAllTables, docTable, firstRow, a, b, c, d, e, columnNum, firstThemeI, lenDocTableCols, lenDocTableRows
      if mainI>0:
         if not isAllTables:
            print('***\nПрограмма нашла ещё одну таблицу. Добавить её в документ? 1 — да, 2 — нет')
            if int(input())!=1:
               return 1
            print('***\nВы хотите, чтобы программа не спрашивала больше и добавила в документ все таблицы? 1 — да, 2 — нет')
            if int(input())==1:
               isAllTables=True
      docTable = doc.tables[mainI]
      firstRow = docTable.row_cells(0)
      a = docTable.row_cells(0)
      b = docTable.row_cells(1)
      c = docTable.row_cells(2)
      d = docTable.row_cells(3)
      e = docTable.row_cells(4)
      columnNum = findColumn('Тема урока', firstRow)
      calculateWidthOfTheme(firstRow, columnNum)
      print(columnNum+1, " — номер колонки, в которой находятся темы")
      firstThemeI = 0
      while docTable.column_cells(columnNum)[firstThemeI].text==docTable.column_cells(columnNum)[0].text:
         firstThemeI+=1 #корректировка длинной (вертикаль) колонки тема урока
      while docTable.column_cells(columnNum)[firstThemeI]._tc==docTable.column_cells(columnNum+widthOfTheme)[firstThemeI]._tc:
         firstThemeI+=1 #корректировка если широкая полоса есть бесполезная
      #firstThemeI = 2
      if (mainI > 0):
         print("Продолжить с первой ячейки новой таблицы, или же первая ячейка - это заголовок и надо отступить?")
         print("1 - продолжить без разрывов; 2 - отступить одну, так как первая есть заголовок")
         if (int(input()) == 1):
            firstThemeI = 0
      i = firstThemeI
      lenDocTableRows = len(docTable.rows)
      lenDocTableCols = len(docTable.columns)
      return 0
   
   def isThereADate():
      k = 0
      global widthOfDate, docTable, columnDate, i, thisCellDate
      while k < widthOfDate:
         if re.findall(r'\d{1,2}[.]\d{1,2}', ''.join(docTable.column_cells(columnDate + k)[i].text.split())):
            columnDate += k
            thisCellDate = docTable.column_cells(columnDate)[i]
            widthOfDate -= k
            return True
         k += 1
      return False
   
   def printProgress():
      global i, lenDocTableRows
      if i%10==0:
         print('Завершено ', i*100//lenDocTableRows, '%')

   def plusOneHourToCell(cell):
      cell.value=cell.value+1
      cell.fill = PatternFill(start_color="00B0F0", end_color="00B0F0", fill_type="solid")
   
   excelI = 2 #в екселе используются человеческие индексы почему-то, не с нуля
   isAllTables=False
   usePlanCell=False
   arrOfHardExcelCells = []
   arrOfHardCellsDates = []
   arrOfHardCellsTexts = []
   arrOfEmptyDates = []
   arrOfEmptyDatesTexts = []
   print('Есть ли в РП дата, или её не стоит учитывать? Если есть - 1, нет - 2')
   if int(input())==1:
      iForAllTables = 0
      for mainI in range(len(doc.tables)):
         if(head() == 1): continue
         columnDate = findColumn('Дата', firstRow)
         widthOfDate = 1
         while columnDate+widthOfDate<len(firstRow) and firstRow[columnDate+widthOfDate]._tc==firstRow[columnDate]._tc:
            widthOfDate+=1
         i=1
         while docTable.cell(i, columnDate)._tc==docTable.cell(i + 1, columnDate+widthOfDate-1)._tc:#для объед яч по вертикали
            i+=1
         if widthOfDate>1 and docTable.cell(i, columnDate)._tc!=docTable.cell(i, columnDate+widthOfDate-1)._tc:
            print('***\nИспользовать для даты ячейку \"'+docTable.cell(i, columnDate).text+'\" или использовать \"'+docTable.cell(i, columnDate+widthOfDate-1).text+'\"?')
            if not usePlanCell:
               print('Введите 1 - если первую, 2 - если вторую (рекомендуется использовать ту, где проставлено больше дат)')
               if int(input())==2:
                  columnDate=columnDate+widthOfDate-1
               if mainI>0:
                  print('Вы хотите, чтобы программа всегда использовала такую ячейку и не спрашивала больше? 1 - да, 2 - нет')
                  if int(input())==1:
                     usePlanCell=True
         #модуль корректировки дат, если ячейки с датами не совпадают с план факт
         notSaveColumnDate=columnDate
         if widthOfDate>2:
            while len(re.findall(r'\d{1,2}[.]\d{1,2}', docTable.column_cells(notSaveColumnDate)[i].text))<1 and notSaveColumnDate<len(docTable.row_cells(i)):
               notSaveColumnDate+=1
            if notSaveColumnDate<len(docTable.row_cells(i)):
               if len(re.findall(r'\d{1,2}[.]\d{1,2}', docTable.column_cells(notSaveColumnDate)[i].text))>=1:
                  columnDate=notSaveColumnDate
         #модуль закончен
         print(columnDate+1, " — номер колонки, в которой находятся даты")
         i = firstThemeI
         #######################################
         while i < lenDocTableRows:
            tempRow=docTable.row_cells(i)
            thisCell = docTable.column_cells(columnNum)[i]
            thisCellDate = docTable.column_cells(columnDate)[i]
            print('Приступили к ячейке ', i + iForAllTables, ', текст:')
            print("---------------------------------------------------------------------")
            print(thisCell.text)
            print("---------------------------------------------------------------------")
            if (thisCell._tc!=docTable.column_cells(columnNum+widthOfTheme)[i]._tc) or isThereADate():#если ячейка не является широкой объед и не является вертикальной
               #оказывается в самой таблице ячейка с темами может идти не ровно, быть 3, а стать 4.
               #напомню, что находимся в левой
               if (thisCell._tc == docTable.column_cells(columnNum)[i - 1]._tc): #если это объединённая ячейка, и по датам тоже
                  if (thisCellDate._tc==docTable.column_cells(columnDate)[i - 1]._tc):
                     i+=1
                     continue
               thisCell = correctThemeHorizontal(thisCell)
               masDates = re.findall(r'\d{1,2}[.]\d{1,2}', ''.join(thisCellDate.text.split()))
               if len(masDates)==1:#если в датах одна дата
                  if masDates[0] not in ''.join(docTable.column_cells(columnDate)[i-1].text.split()):#если дата не равна дате сверху
                     appendToExcelWithoutRepeat(thisCell.text)
                  else:
                     plusOneHourToCell(sheet.cell(row=excelI-1, column=5))
                     appendToExcelWithMerge(thisCell.text, docTable.column_cells(columnNum)[i-1].text)   
               elif len(masDates) == 0:
                  if not isThereADate():
                     print('!!!!!!!!!!!!!!!!!!!!!!!')
                     print('Внимание, обнаружена ячейка с пустой датой. В таком случае программа не добавляет урок! (так как скорее всего он уже добавлен, но нужно перепроверить потом) Запомните место и отредактируйте в excel самостоятельно')
                     print('Нажмите \"Ввод\", чтобы продолжить')
                     arrOfEmptyDates.append(i + iForAllTables)
                     arrOfEmptyDatesTexts.append(thisCell.text)
                     input()
                  else:
                     continue
               else:
                  #в датах несколько дат!
                  #посчитаем, объединённая эта ячейка или одна
                  cellLen = 1#колво ячеек с темами, если яч с датами общая
                  j=i
                  while j<len(docTable.rows)-1:
                     if docTable.column_cells(columnDate)[j]._tc==docTable.column_cells(columnDate)[j+1]._tc:
                        #экспериментальный иф, о проблеме объединённых ячеек дат и объединённых ячеек тем (это всё равно считаем одна большая яч)
                        if docTable.column_cells(columnNum)[j]._tc!=docTable.column_cells(columnNum)[j+1]._tc:
                           cellLen+=1
                     else:
                        break
                     j+=1
                  cellCount = 0#итератор по ячейкам с темами, но в первом случае - по темам через энтер
                  #если одна, то сравниваем сначала с пред потом внутри
                  if cellLen == 1:
                     #напомню, случай, когда одна яч и там, и там, а даты и темы прописаны через энтер (одна яч тема одна яч дата, без объединений)
                     masThemes = re.split('\n+', thisCell.text)
                     lenMasThemes = len(masThemes)
                     lenMasDates = len(masDates) #>=2
                     ohapka = 1
                     dateIndex = 0
                     if (lenMasThemes >= lenMasDates):
                        ohapka = lenMasThemes / lenMasDates
                        if (ohapka % 1 > 0.5):#7 на 3 это 2 2 и 3, а 8 на 3 - это 3 3 и 2
                           ohapka += 1
                        ohapka = int(ohapka)
                     # c пред яч
                     takeAllFlag = 0
                     if (lenMasThemes > 1):
                        print("Обнаружено несколько строк в ячейке и несколько дат. Постараться распределить строки на разные числа или брать все строки и создать столько одинаковых уроков, сколько дат?")
                        print("Короче разбивать на несколько тем (сколько дат - столько тем) или не разбивать?")
                        print("Нажмите 1 или 2")
                        if (int(input()) == 2):
                           ohapka = lenMasThemes
                           takeAllFlag = 1
                        else:
                           arrOfHardCellsDates.append(docTable.column_cells(columnDate)[i].text)
                     
                     gran = cellCount + ohapka
                     pastDateIndex = docTable.column_cells(columnDate)[i-1].text
                     while (cellCount < gran) and (cellCount < lenMasThemes):
                        if masDates[dateIndex] not in pastDateIndex:
                           if (takeAllFlag == 0): arrOfHardExcelCells.append(excelI - 1)
                           appendToExcelWithoutRepeat(masThemes[cellCount])
                        else:
                           if (masDates[dateIndex] != pastDateIndex):
                              plusOneHourToCell(sheet.cell(row=excelI-1, column=5))
                           appendToExcelWithMerge(masThemes[cellCount], docTable.column_cells(columnNum)[i-1].text)
                        cellCount += 1
                        pastDateIndex = masDates[dateIndex]
                     j = cellCount
                     dateIndex += 1
                     if (takeAllFlag): j = 0
                     #теперь внутри
                     #попытаемся равномерно распределить темы по датам (строки через \n), ведь их может быть 3, а дат 2 например
                     while j < max(lenMasThemes, lenMasDates):
                        cellCount = j
                        if cellCount >= lenMasThemes:
                           cellCount = lenMasThemes - 1
                        pastDateIndex = dateIndex - 1
                        gran = cellCount + ohapka
                        if (dateIndex == lenMasDates - 1):
                           gran = lenMasThemes#идёт до конца массива, а не охапки в последнем случае
                        while (cellCount < gran) and (cellCount < lenMasThemes):
                           if masDates[dateIndex] != masDates[pastDateIndex]:
                              if (takeAllFlag == 0): arrOfHardExcelCells.append(excelI - 1)
                              appendToExcelWithoutRepeat(masThemes[cellCount])
                           else:
                              if (dateIndex != pastDateIndex):
                                 plusOneHourToCell(sheet.cell(row=excelI-1, column=5))
                              appendToExcelWithMerge(masThemes[cellCount], masThemes[cellCount-1])
                           cellCount += 1
                           pastDateIndex = dateIndex
                        j = cellCount
                        if (takeAllFlag): 
                           j = 0
                           if (dateIndex == lenMasDates - 1): break
                        dateIndex += 1
                  #теперь несколько строк, но яч с датами одна (дат много)
                  else:
                     arrOfHardCellsDates.append(docTable.column_cells(columnDate)[i].text)
                     if masDates[cellCount] not in docTable.column_cells(columnDate)[i-1].text:
                        arrOfHardExcelCells.append(excelI - 1)
                        appendToExcelWithoutRepeat(thisCell.text)
                     else:
                        plusOneHourToCell(sheet.cell(row=excelI-1, column=5))
                        appendToExcelWithMerge(thisCell.text, docTable.column_cells(columnNum)[i-1].text)
                           
                     for cellCount in range(1, cellLen):
                        if cellCount >= len(masDates):
                           #если получилось так, что строк больше чем дат, то тупо крепим всё к предыдущей
                           #ведь cellLen - колво строк, а masDates - колво дат
                           sheet.cell(row=excelI-1, column=2).value=sheet.cell(row=excelI-1, column=2).value+' '+" ".join(docTable.column_cells(columnNum)[i+cellCount].text.split())
                        else:
                           if masDates[cellCount]!=masDates[cellCount-1]:
                              arrOfHardExcelCells.append(excelI - 1)
                              appendToExcelWithoutRepeat(docTable.column_cells(columnNum)[i+cellCount].text)
                           else:
                              plusOneHourToCell(sheet.cell(row=excelI-1, column=5))
                              appendToExcelWithMerge(docTable.column_cells(columnNum)[i+cellCount].text, docTable.column_cells(columnNum)[i+cellCount-1].text)
                     i+=cellCount
            else:
               arrOfHardCellsTexts.append(thisCell.text)
            i+=1
            printProgress()
         iForAllTables += i
   else:
      iForAllTables = 0
      for mainI in range(len(doc.tables)):
         if(head() == 1): continue
         while i < lenDocTableRows:
            tempRow=docTable.row_cells(i)
            thisCell = docTable.column_cells(columnNum)[i]
            print('Приступили к ячейке ', i + iForAllTables)
            print('Текст: ', thisCell.text)
            if thisCell._tc!=docTable.column_cells(columnNum+widthOfTheme)[i]._tc:#если ячейка не является широкой объед и не является вертикальной
               #оказывается в самой таблице ячейка с темами может идти не ровно, быть 3, а стать 4.
               thisCell = correctThemeHorizontal(thisCell)
               if i!=lenDocTableRows-1:#защита индекса
                  while docTable.column_cells(columnNum)[i+1]._tc==docTable.column_cells(columnNum)[i]._tc and i < lenDocTableRows-1:#корректировка для объед по вертикали, так как мы без дат
                     i+=1
               appendToExcelWithoutRepeat(docTable.column_cells(columnNum)[i].text)
            i+=1
            printProgress()
         iForAllTables += i
   print('Прежде чем приступить к созданию excel файла, убедитесь, что он сейчас не используется и закрыт. Нажмите \"Ввод\" для продолжения.')
   input()
   if nameConfFile in os.listdir():
      config.read(nameConfFile)
      initialdir = config.get("Main", "savepath")
   else:
      open(nameConfFile, 'x')
      config.add_section("Main")
      initialdir = os.getenv('USERPROFILE')+"\\Desktop"
   nameExcelFile = tkinter.filedialog.asksaveasfile(initialdir = initialdir, title = "Выберите, куда сохранить excel файл", filetypes = [('Только excel файлы', '.xlsx .xls')], defaultextension=".xlsx").name
   wb.save(nameExcelFile)
   
   if nameExcelFile:
      config.set("Main", "savepath", "/".join(nameExcelFile.split("/")[0:-1]))
      with open(nameConfFile, 'w') as configFile:
         config.write(configFile)
   print('Успех')
   if arrOfEmptyDates:
      print('Были обнаружены пустые даты, вот приблизительные номера этих строк:')
      print(arrOfEmptyDates)
      print("Тексты этих ячеек с пустыми датами:")
      print(arrOfEmptyDatesTexts)
   if arrOfHardExcelCells:
      print('Внимание! Возможна склейка дат либо неправильное распределение тем в следующих ячейках (нумерация тем в excel документе, для уточнения смотрите даты):')
      print(arrOfHardExcelCells)
      print('Вот их даты:')
      print(arrOfHardCellsDates)
   if arrOfHardCellsTexts:
      print("Внимание! Не добавлены следующие строки:")
      print(arrOfHardCellsTexts)
   print('Обязательно проверьте ячейки, в которых программа выставила 2 и более часов')
   print('Программа завершила своё выполнение, нажмите \"Ввод\" дважды, чтобы выйти.')
   input()
   input()
else:
   print('Файл не выбран, ошибка')
   input()
